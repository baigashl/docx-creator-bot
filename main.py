import json
import os
from aiogram import types
from aiogram.dispatcher.filters import Text
from aiogram.utils import executor
from aiogram.types import ReplyKeyboardMarkup, KeyboardButton
import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from config import dp, bot
from database.database import Order
from get_data import get_delivery_data_from_api, get_pickup_data_from_api
from docx.shared import Pt, Cm, Inches, Mm
import docx2txt

button1 = KeyboardButton('🚚 Доставка')
button2 = KeyboardButton('Самовывоз')
markup3 = ReplyKeyboardMarkup(resize_keyboard=True).add(button1, button2)


@dp.message_handler(commands=['start'])
async def process_start_command(message: types.Message):
    await bot.send_message(message.from_user.id, "👋🏼Добро пожаловать", reply_markup=markup3)


@dp.message_handler(Text(equals='🚚 Доставка'))
async def add_db(message: types.Message):
    id_user = message.from_user.id
    await bot.send_message(id_user, "Немного подождите 😊")
    await get_delivery_data_from_api()
    with open('delivery.json') as json_file:
        delivery_data = json.load(json_file)

    doc = docx.Document()
    sections = doc.sections
    section = sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    # section.header_distance = Mm(12.7)
    # section.footer_distance = Mm(12.7)

    font_styles = doc.styles
    font_charstyle = font_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
    font_object = font_charstyle.font
    font_object.size = Pt(9)
    font_object.name = 'Helvetica'
    for section in sections:
        section.top_margin = Inches(0.2)
        section.bottom_margin = Inches(0)
        section.left_margin = Inches(2)
        section.right_margin = Inches(2)
    try:
        # print(delivery_data[0]["orders"][0]["orderCode"])
        for order in delivery_data[0]["orders"]:
            code = Order.get_or_none(orderCode=order["orderCode"])
            if code is None:

                Order.create(
                    orderCode=order["orderCode"],
                )
                with open(f'{order["orderCode"]}_prod_detail.json') as json_file:
                    detail_data = json.load(json_file)

                paragraph = doc.add_paragraph(f'Номер заказа\n')
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                paragraph.add_run(f'{detail_data["orderId"]}\n', style='CommentsStyle')
                paragraph.add_run(f'Сумма\n', style='CommentsStyle').bold = True
                paragraph.add_run(f'{detail_data["localizedSum"]}\n', style='CommentsStyle')
                paragraph.add_run('\n', style='CommentsStyle')
                paragraph.add_run(f'Стоимость доставки для клиента\n', style='CommentsStyle')
                paragraph.add_run(f'{detail_data["deliveryDiscountFormatted"]}\n', style='CommentsStyle')
                paragraph.add_run('\n', style='CommentsStyle')
                paragraph.add_run(f'Имя\n', style='CommentsStyle')
                paragraph.add_run(f'{detail_data["purchaserFirstName"]} ', style='CommentsStyle')
                paragraph.add_run(f'{detail_data["purchaserLastName"]}\n', style='CommentsStyle')
                paragraph.add_run(f'Номер телефона\n', style='CommentsStyle')
                paragraph.add_run(f'+7 ({detail_data["purchaserPhoneNumber"][0:3]})-{detail_data["purchaserPhoneNumber"][3:6]}-{detail_data["purchaserPhoneNumber"][6:8]}-{detail_data["purchaserPhoneNumber"][8:10]}\n', style='CommentsStyle')
                paragraph.add_run('\n', style='CommentsStyle')
                #################################################################3
                for product in detail_data["products"]:
                    paragraph.add_run(f'{product["name"]}\n', style='CommentsStyle')
                    paragraph.add_run(f'{product["quantity"]}                 ', style='CommentsStyle')
                    paragraph.add_run(f'{product["localizedActualPrice"]}\n', style='CommentsStyle')
                ########################################################
                # paragraph.add_run(f'{order["entries"][0]["name"]}\n', style='CommentsStyle')
                # paragraph.add_run(f'{order["entries"][0]["quantity"]}                 ', style='CommentsStyle')
                # paragraph.add_run(f'{order["entries"][0]["totalPrice"]}\n', style='CommentsStyle')
                ###############################################################################3
                paragraph.add_run('\n', style='CommentsStyle')
                paragraph.add_run(f'Адрес доставки\n', style='CommentsStyle').bold = True
                paragraph.add_run('\n', style='CommentsStyle')
                paragraph.add_run(f'{detail_data["deliveryAddress"]["formattedAddress"]}\n', style='CommentsStyle')
                doc.add_page_break()
                # doc.center()

                os.remove(f'{order["orderCode"]}_prod_detail.json')
            else:
                os.remove(f'{order["orderCode"]}_prod_detail.json')
                print('exist')

        doc.save(f'delivery.docx')

        my_text = docx2txt.process("delivery.docx")
        if len(my_text) < 1:
            await bot.send_message(id_user, "На данный момент нет НОВЫХ доставок 🙃")
        else:
            await bot.send_document(id_user, document=open(f'delivery.docx', 'rb'))

        os.remove(f'delivery.docx')
        # await bot.send_file()
        await bot.send_message(id_user, "Все данные были добавлены в Базу данных 😉 👌")
    except KeyError:
        await bot.send_message(id_user, "На данный момент нет доставок 🙃")



@dp.message_handler(Text(equals='Самовывоз'))
async def pickup(message: types.Message):
    id_user = message.from_user.id
    await bot.send_message(id_user, "Немного подождите 😊")
    await get_pickup_data_from_api()
    with open('pickup.json') as json_file:
        delivery_data = json.load(json_file)

    doc = docx.Document()
    sections = doc.sections
    section = sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    # section.header_distance = Mm(12.7)
    # section.footer_distance = Mm(12.7)

    font_styles = doc.styles
    font_charstyle = font_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
    font_object = font_charstyle.font
    font_object.size = Pt(9)
    font_object.name = 'Helvetica'
    for section in sections:
        section.top_margin = Inches(0.2)
        section.bottom_margin = Inches(0)
        section.left_margin = Inches(2)
        section.right_margin = Inches(2)
    try:
        # print(delivery_data[0]["orders"][0]["orderCode"])
        for order in delivery_data[0]["orders"]:
            code = Order.get_or_none(orderCode=order["orderCode"])
            if code is None:

                Order.create(
                    orderCode=order["orderCode"],
                )
                with open(f'{order["orderCode"]}_pickup_prod_detail.json') as json_file:
                    detail_data = json.load(json_file)

                paragraph = doc.add_paragraph(f'Номер заказа\n')
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                paragraph.add_run(f'{detail_data["orderId"]}\n', style='CommentsStyle')
                paragraph.add_run(f'Сумма\n', style='CommentsStyle').bold = True
                paragraph.add_run(f'{detail_data["localizedSum"]}\n', style='CommentsStyle')
                paragraph.add_run('\n', style='CommentsStyle')
                paragraph.add_run(f'Стоимость доставки для клиента\n', style='CommentsStyle')
                paragraph.add_run(f'{detail_data["deliveryDiscountFormatted"]}\n', style='CommentsStyle')
                paragraph.add_run('\n', style='CommentsStyle')
                paragraph.add_run(f'Имя\n', style='CommentsStyle')
                paragraph.add_run(f'{detail_data["purchaserFirstName"]} ', style='CommentsStyle')
                paragraph.add_run(f'{detail_data["purchaserLastName"]}\n', style='CommentsStyle')
                paragraph.add_run(f'Номер телефона\n', style='CommentsStyle')
                paragraph.add_run(f'+7 ({detail_data["purchaserPhoneNumber"][0:3]})-{detail_data["purchaserPhoneNumber"][3:6]}-{detail_data["purchaserPhoneNumber"][6:8]}-{detail_data["purchaserPhoneNumber"][8:10]}\n', style='CommentsStyle')
                paragraph.add_run('\n', style='CommentsStyle')
                #################################################################3
                for product in detail_data["products"]:
                    paragraph.add_run(f'{product["name"]}\n', style='CommentsStyle')
                    paragraph.add_run(f'{product["quantity"]}                 ', style='CommentsStyle')
                    paragraph.add_run(f'{product["localizedActualPrice"]}\n', style='CommentsStyle')
                ########################################################
                # paragraph.add_run(f'{order["entries"][0]["name"]}\n', style='CommentsStyle')
                # paragraph.add_run(f'{order["entries"][0]["quantity"]}                 ', style='CommentsStyle')
                # paragraph.add_run(f'{order["entries"][0]["totalPrice"]}\n', style='CommentsStyle')
                ###############################################################################3
                doc.add_page_break()
                # doc.center()

                os.remove(f'{order["orderCode"]}_pickup_prod_detail.json')
            else:
                os.remove(f'{order["orderCode"]}_pickup_prod_detail.json')
                print('exist')

        doc.save(f'pickup.docx')

        my_text = docx2txt.process("pickup.docx")
        if len(my_text) < 1:
            await bot.send_message(id_user, "На данный момент нет НОВЫХ Самовывозов 🙃")
        else:
            await bot.send_document(id_user, document=open(f'pickup.docx', 'rb'))

        os.remove(f'pickup.docx')
        # await bot.send_file()
        await bot.send_message(id_user, "Все данные были добавлены в Базу данных 😉 👌")
    except KeyError:
        await bot.send_message(id_user, "На данный момент нет Самовывоза 🙃")



if __name__ == '__main__':
    executor.start_polling(dp)
